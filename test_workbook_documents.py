import unittest
from io import BytesIO

from docx import Document

from file_utils import export_instructor_guide_to_word, export_student_workbook_to_word


class WorkbookDocumentsTest(unittest.TestCase):
    ITEMS = [
        {
            "unit_code": "ICT/301",
            "unit_title": "Install Network Equipment",
            "lo_num": "1",
            "lo_description": "Install network devices",
            "pc_code": "1.1",
            "pc_description": "Install a switch",
            "question_type": "Scenario-Based",
            "question": "Describe how to install the network switch safely.",
            "weight": 5,
            "ideal_answer": ["Isolate power and follow the manufacturer instructions."],
            "marking_scheme": ["Safe installation: 5 marks"],
        },
        {
            "unit_code": "ICT/302",
            "unit_title": "Test Network Services",
            "lo_num": "1",
            "lo_description": "Test connections",
            "pc_code": "1.2",
            "pc_description": "Test connectivity",
            "question_type": "Direct",
            "question": "Which test confirms network connectivity?",
            "weight": 5,
            "ideal_answer": ["Use an appropriate connectivity test."],
            "marking_scheme": ["Correct test: 5 marks"],
        },
    ]

    def test_student_document_layout_and_content(self):
        document = Document(BytesIO(export_student_workbook_to_word("ICT", "Level 3", "Ada", self.ITEMS)))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertLess(text.index("Trade Name:"), text.index("Unit ICT/301:"))
        self.assertIn("Student Name: Ada", text)
        self.assertNotIn("Ideal Answer:", text)
        self.assertEqual(text.count("Answer:"), 2)
        self.assertEqual(text.count("________________________________________________________________________"), 16)
        self.assertGreaterEqual(document.element.body.xml.count('w:type="page"'), 2)

    def test_documents_share_exact_question_text(self):
        student = Document(BytesIO(export_student_workbook_to_word("ICT", "Level 3", "Ada", self.ITEMS)))
        instructor = Document(BytesIO(export_instructor_guide_to_word("ICT", "Level 3", "Ada", self.ITEMS)))
        student_text = "\n".join(paragraph.text for paragraph in student.paragraphs)
        instructor_text = "\n".join(paragraph.text for paragraph in instructor.paragraphs)
        for item in self.ITEMS:
            self.assertEqual(student_text.count(item["question"]), 1)
            self.assertEqual(instructor_text.count(item["question"]), 1)
        self.assertIn("Ideal Answer:", instructor_text)
        self.assertIn("Marking Scheme:", instructor_text)


if __name__ == "__main__":
    unittest.main()
