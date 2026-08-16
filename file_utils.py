from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
import re
import zipfile

MAPPING_UNIT_CODE_PATTERN = r"[A-Za-z0-9/._-]+"

def set_cell_border(cell, **kwargs):
    """
    Utility to set specific cell borders using XML.
    Example: set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_background(cell, color):
    """Set cell background color (hex string like 'D9D9D9')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_page_number(run):
    """Adds a dynamic page number field to a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def get_unit_number(unit_code):
    """Extract the 3-digit unit number from supported unit code formats."""
    try:
        parts = [part.strip() for part in str(unit_code).split("/") if part.strip()]
        for part in parts:
            if re.fullmatch(r"\d{3}", part):
                # Convert to int then back to string to remove leading zeros (e.g., '008' -> '8')
                return str(int(part))
        return str(unit_code).strip()
    except (ValueError, IndexError, AttributeError):
        return str(unit_code).strip()

def _parse_single_mapping(inner_content):
    # Split the inner content by semicolon, but only if a new Unit Code follows.
    segments = re.split(rf';\s*(?={MAPPING_UNIT_CODE_PATTERN}\s*-)', inner_content)
    
    unit_nums = []
    mapping_parts = []

    for segment in segments:
        if '-' in segment:
            u_code, mapping = segment.split('-', 1)
            unit_nums.append(get_unit_number(u_code.strip()))
            mapping_parts.append(mapping.strip())
        else:
            mapping_parts.append(segment.strip())
            
    final_units = ", ".join(dict.fromkeys(unit_nums))
    final_mapping = "; ".join(mapping_parts)
    return final_units, final_mapping

def parse_report_chunks(text):
    """
    Parses the full text and yields chunks of (narrative, unit, mapping, is_last_in_paragraph).
    """
    raw_lines = [line.strip() for line in text.splitlines() if line.strip()]
    merged_lines = []
    for line in raw_lines:
        if re.match(rf'^\({MAPPING_UNIT_CODE_PATTERN}\s*-\s*(?:LO)?\s*\d+', line, flags=re.IGNORECASE) and merged_lines:
            merged_lines[-1] = f"{merged_lines[-1]} {line}"
        else:
            merged_lines.append(line)
            
    chunks = []
    
    for line in merged_lines:
        parts = re.split(rf'(\({MAPPING_UNIT_CODE_PATTERN}\s*-\s*(?:LO)?\s*\d+[^)]*\))', line, flags=re.IGNORECASE)
        
        line_chunks = []
        narrative_acc = ""
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            if re.match(rf'^\({MAPPING_UNIT_CODE_PATTERN}\s*-\s*(?:LO)?\s*\d+[^)]*\)$', part, flags=re.IGNORECASE):
                mapping_str = part[1:-1]
                u_num, mapping = _parse_single_mapping(mapping_str)
                # Clean up any trailing punctuation on the narrative
                cleaned_narrative = re.sub(r'\s*[.,;:]+$', '', narrative_acc.strip()).strip()
                line_chunks.append({
                    'narrative': cleaned_narrative,
                    'unit': u_num,
                    'mapping': mapping
                })
                narrative_acc = ""
            else:
                narrative_acc += (" " if narrative_acc else "") + part
                
        if narrative_acc.strip():
            cleaned_narrative = re.sub(r'\s*[.,;:]+$', '', narrative_acc.strip()).strip()
            line_chunks.append({
                'narrative': cleaned_narrative,
                'unit': "",
                'mapping': ""
            })
            
        for i, lc in enumerate(line_chunks):
            lc['is_last_in_paragraph'] = (i == len(line_chunks) - 1)
            chunks.append(lc)
            
    return chunks

def _create_official_nsq_template(doc, candidate_name, date, report_text, units_summary, criteria_summary, evidence_type="observation"):
    """
    Internal helper to generate the official CPN-ARF-02 Performance Evidence Record Form.
    """
    # Global Style Adjustments
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set Page Margins to 0.75"
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- ADD PAGE NUMBERS (Footer) ---
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Page ")
    add_page_number(footer_run)

    # --- HEADER SECTION (Table 1) ---
    header_table = doc.add_table(rows=1, cols=1)
    header_table.width = Inches(7.0)
    rowhead = header_table.rows[0]
    rowhead.height = Inches(0.6) # 1. Make header row taller
    header_table.style = 'Table Grid'
    
    header_cell = header_table.cell(0, 0)
    header_cell.text = "Ref:ARF02A      NATIONAL SKILLS QUALIFICATION"
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Apply font size and bold to the text in the cell
    for paragraph in header_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(14)
            run.bold = True

    # --- UNDERLINED TITLE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nASSESSORS AND CANDIDATES PERFORMANCE EVIDENCE RECORD FORM")
    run.bold = True
    run.underline = True
    run.size = Pt(12)

    # --- USER METADATA (Table 2) ---
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = 'Table Grid'
    meta_table.width = Inches(7.0)
    rowmeta = meta_table.rows[0]
    rowmeta.height = Inches(0.4) # 2. Make metadata row taller
    
    cell_candidate = meta_table.cell(0, 0)
    cell_candidate.text = f"Candidate Name: {candidate_name}"
    cell_candidate.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_units = meta_table.cell(0, 1)
    cell_units.text = f"Units: {units_summary}"
    cell_units.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    # --- INSTRUCTION BLOCK ---
    instr = doc.add_paragraph("\nThis form can be used for Assessor’s observation of practical workplace activities or Candidates statement of practical activities. NB: Assessor may wish to ask a candidate some questions relating to the activity and record the question and the answer in this form also.")
    instr.italic = True
    instr.paragraph_format.space_after = Pt(10)

    # Clean report text: Remove summary block and split into paragraphs
    clean_report = report_text.split("----- SUMMARY OF CRITERIA COVERED -----")[0].strip()

    # Parse the clean report into chunks (sentence + mapping pairs)
    chunks = parse_report_chunks(clean_report)

    # --- MAIN ACTIVITY GRID (Table 3) ---
    # We create a table with 3 header rows + one row for every chunk
    num_data_rows = max(1, len(chunks))
    main_table = doc.add_table(rows=3 + num_data_rows, cols=5)
    main_table.style = 'Table Grid'
    main_table.autofit = False
    main_table.width = Inches(7.0)

    # Force fixed table layout at XML level to ensure widths are strictly respected
    tbl = main_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # Set column widths (Unit: 0.58, LO: 1.50, Sub-columns: 1.64 each)
    # Total width sums to exactly 7.0" to fit perfectly within 0.75" margins
    col_widths = [0.58, 1.50, 1.64, 1.64, 1.64]
    for i, width in enumerate(col_widths):
        main_table.columns[i].width = Inches(width)
        for cell in main_table.columns[i].cells:
            cell.width = Inches(width)

    # Header Row 1: Merging for "Tick which evidence method"
    main_table.cell(0, 0).text = "Unit"
    main_table.cell(0, 1).text = "LO/Assessment criteria"
    tick_header = main_table.cell(0, 2).merge(main_table.cell(0, 4))
    tick_header.text = "Tick which evidence method"
    tick_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header Row 2 & 3: Merging Unit and LO vertically across header rows
    main_table.cell(0, 0).merge(main_table.cell(2, 0))
    main_table.cell(0, 1).merge(main_table.cell(2, 1))

    # Evidence Checkboxes
    check_obs = "☑" if evidence_type == "observation" else "☐"
    check_wit = "☑" if evidence_type == "witness" else "☐"
    check_per = "☑" if evidence_type == "personal" else "☐"
    
    main_table.cell(1, 2).text = f"{check_obs} Observation"
    main_table.cell(1, 3).text = f"{check_wit} Witness statement"
    main_table.cell(1, 4).text = f"{check_per} Personal statement"

    # Record Label
    record_label = main_table.cell(2, 2).merge(main_table.cell(2, 4))
    record_label.text = "Record of observed activity and performance."
    record_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- STYLE HEADER ROWS (Bold + Shading) ---
    for row_idx in range(3):
        for cell in main_table.rows[row_idx].cells:
            set_cell_background(cell, 'EFEFEF') # Light gray shading
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    # --- DATA ROWS ---
    if not chunks:
        row = main_table.rows[3]
        row.cells[0].text = units_summary
        row.cells[1].text = criteria_summary
        row.cells[2].merge(row.cells[4]).text = clean_report
    else:
        for i, chunk in enumerate(chunks):
            row = main_table.rows[3 + i]
            
            row.cells[0].text = chunk['unit']
            row.cells[1].text = chunk['mapping']
            
            # 6. Dynamic Font Sizing for long LOs
            if len(chunk['mapping']) > 25:
                for p in row.cells[1].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
            
            # Merge the narrative cells and insert text
            content_cell = row.cells[2].merge(row.cells[4])
            content_cell.text = chunk['narrative']

            # Add space after each entry. 
            # If it's the last chunk of a paragraph, use 12pt space. Otherwise 6pt for connected sentences.
            space = Pt(12) if chunk['is_last_in_paragraph'] else Pt(6)
            row.cells[0].paragraphs[0].paragraph_format.space_after = space
            row.cells[1].paragraphs[0].paragraph_format.space_after = space
            content_cell.paragraphs[0].paragraph_format.space_after = space

            # Adjust borders to make multiple rows look like one continuous box
            # We remove the internal horizontal lines between chunks
            if len(chunks) > 1:
                # If not the last chunk, remove bottom border
                if i < len(chunks) - 1:
                    for cell in row.cells:
                        set_cell_border(cell, bottom={"val": "nil"})
                # If not the first chunk, remove top border
                if i > 0:
                    for cell in row.cells:
                        set_cell_border(cell, top={"val": "nil"})

    # Apply top vertical alignment to all data rows
    for row in main_table.rows[3:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # --- DECLARATION ---
    doc.add_paragraph("\nI confirm that the evidence listed above is true record of the performed activities observed during this assessment.")

    # --- SIGNATURES GRID (Table 4) ---
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'
    
    sigs = [
        "Work base Assessor / witness Signature:",
        "Candidate's Signature:",
        "Assessor's Signature:",
        "Internal Verifier's Signature:"
    ]
    
    for i, label in enumerate(sigs):
        row = sig_table.rows[i]
        row.height = Inches(0.6) # 3. Make signature rows taller
        
        # Add labels and vertical alignment
        cell_left = sig_table.cell(i, 0)
        cell_right = sig_table.cell(i, 1)
        
        cell_left.text = f"{label} __________________"
        cell_right.text = f"Date: {date}"
        
        cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def export_to_word(name, date, report_text, assessor_name, timeline="N/A", atmosphere="N/A", selected_pcs=None):
    """
    Generates a standardized NSQ Word report.
    """
    doc = Document()
    
    units_list = []
    criteria_list = []

    if selected_pcs and isinstance(selected_pcs, list):
        for pc_str in selected_pcs:
            parts = pc_str.split(' - ')
            if len(parts) >= 3:
                units_list.append(parts[0].strip())
                criteria_list.append(f"LO{parts[1]}:{parts[2].split(':')[0]}")
    elif isinstance(selected_pcs, str) and selected_pcs:
        units_list = [u.strip() for u in selected_pcs.split(',') if u.strip()]

    units_summary = ", ".join(dict.fromkeys([get_unit_number(u) for u in units_list]))
    criteria_summary = "; ".join(criteria_list)

    _create_official_nsq_template(
        doc=doc,
        candidate_name=name,
        date=date,
        report_text=report_text,
        units_summary=units_summary,
        criteria_summary=criteria_summary,
        evidence_type="observation"
    )
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_witness_to_word(witness_name, witness_role, candidate_name, date, statement_text, selected_pcs=None):
    """
    Generates a standardized NSQ Witness Statement Word document.
    """
    doc = Document()
    
    units_list = []
    criteria_list = []

    if selected_pcs and isinstance(selected_pcs, list):
        for pc_str in selected_pcs:
            parts = pc_str.split(' - ')
            if len(parts) >= 3:
                units_list.append(parts[0].strip())
                criteria_list.append(f"LO{parts[1]}:{parts[2].split(':')[0]}")
    else:
        if isinstance(selected_pcs, str) and selected_pcs:
            units_list = [u.strip() for u in selected_pcs.split(',') if u.strip()]
        criteria_list = ["Refer to narrative"]

    units_summary = ", ".join(dict.fromkeys([get_unit_number(u) for u in units_list]))
    criteria_summary = "; ".join(criteria_list)

    _create_official_nsq_template(
        doc=doc,
        candidate_name=candidate_name,
        date=date,
        report_text=statement_text,
        units_summary=units_summary,
        criteria_summary=criteria_summary,
        evidence_type="witness"
    )

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_personal_statement_to_word(name, date, statement_text, selected_pcs=None):
    """
    Generates a standardized NSQ Personal Statement Word document.
    """
    doc = Document()
    
    units_list = []
    criteria_list = []

    if selected_pcs and isinstance(selected_pcs, list):
        for pc_str in selected_pcs:
            parts = pc_str.split(' - ')
            if len(parts) >= 3:
                units_list.append(parts[0].strip())
                criteria_list.append(f"LO{parts[1]}:{parts[2].split(':')[0]}")
    else:
        if isinstance(selected_pcs, str) and selected_pcs:
            units_list = [u.strip() for u in selected_pcs.split(',') if u.strip()]
        criteria_list = ["Refer to narrative"]

    units_summary = ", ".join(dict.fromkeys([get_unit_number(u) for u in units_list]))
    criteria_summary = "; ".join(criteria_list)

    _create_official_nsq_template(
        doc=doc,
        candidate_name=name,
        date=date,
        report_text=statement_text,
        units_summary=units_summary,
        criteria_summary=criteria_summary,
        evidence_type="personal"
    )

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def _build_nsq_pdf_story(name, date, report_text, assessor_name, units_summary, evidence_type="observation"):
    """
    Internal helper to build the ReportLab story matching the official NSQ form.
    Returns a list of flowables ready for SimpleDocTemplate.build().
    """
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'NsqTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        alignment=1,
        spaceAfter=8
    )
    header_cell_style = ParagraphStyle(
        'NsqHeaderCell',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        alignment=1
    )
    cell_style = ParagraphStyle(
        'NsqCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=11
    )
    cell_bold_style = ParagraphStyle(
        'NsqCellBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11
    )
    cell_center_style = ParagraphStyle(
        'NsqCellCenter',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        alignment=1
    )
    instruction_style = ParagraphStyle(
        'NsqInstruction',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8,
        leading=10,
        spaceAfter=8
    )
    narrative_style = ParagraphStyle(
        'NsqNarrative',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=11
    )

    story = []

    # --- HEADER TABLE (Table 1) ---
    header_table = Table(
        [[Paragraph("Ref:ARF02A      NATIONAL SKILLS QUALIFICATION", header_cell_style)]],
        colWidths=[6.5 * inch]
    )
    header_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(header_table)

    # --- UNDERLINED TITLE ---
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        "<u>ASSESSORS AND CANDIDATES PERFORMANCE EVIDENCE RECORD FORM</u>",
        title_style
    ))

    # --- METADATA TABLE (Table 2) ---
    clean_report = report_text.split("----- SUMMARY OF CRITERIA COVERED -----")[0].strip()
    meta_table = Table(
        [
            [
                Paragraph(f"Candidate Name: {name}", cell_style),
                Paragraph(f"Units: {units_summary}", cell_style)
            ]
        ],
        colWidths=[3.25 * inch, 3.25 * inch]
    )
    meta_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(meta_table)

    # --- INSTRUCTION BLOCK ---
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        "This form can be used for Assessor’s observation of practical workplace activities "
        "or Candidates statement of practical activities. NB: Assessor may wish to ask a "
        "candidate some questions relating to the activity and record the question and the "
        "answer in this form also.",
        instruction_style
    ))

    # --- PARSE REPORT CHUNKS ---
    chunks = parse_report_chunks(clean_report)
    if not chunks:
        chunks = [{'narrative': clean_report, 'unit': '', 'mapping': '', 'is_last_in_paragraph': True}]

    # --- MAIN ACTIVITY GRID (Table 3) ---
    # Columns: Unit, LO/Assessment criteria, Tick which evidence method (merged 3 cols), Record label
    # Word layout: 5 columns (Unit 0.58", LO 1.50", 3 evidence cols 1.64" each)
    # PDF widths: keep proportional
    col_widths = [0.58 * inch, 1.50 * inch, 1.64 * inch, 1.64 * inch, 1.64 * inch]

    main_data = []
    # Header Row 1
    main_data.append([
        Paragraph("Unit", cell_center_style),
        Paragraph("LO/Assessment criteria", cell_center_style),
        Paragraph("Tick which evidence method", cell_center_style),
        "",
        ""
    ])
    # Header Row 2 (evidence checkboxes)
    check_obs = "[X]" if evidence_type == "observation" else "[ ]"
    check_wit = "[X]" if evidence_type == "witness" else "[ ]"
    check_per = "[X]" if evidence_type == "personal" else "[ ]"
    main_data.append([
        "",
        "",
        Paragraph(f"{check_obs} Observation", cell_style),
        Paragraph(f"{check_wit} Witness statement", cell_style),
        Paragraph(f"{check_per} Personal statement", cell_style),
    ])
    # Header Row 3 (record label)
    main_data.append([
        "",
        "",
        Paragraph("Record of observed activity and performance.", cell_center_style),
        "",
        ""
    ])

    # Data rows
    for chunk in chunks:
        main_data.append([
            Paragraph(chunk['unit'], cell_style),
            Paragraph(chunk['mapping'], cell_style),
            Paragraph(chunk['narrative'], narrative_style),
            "",
            ""
        ])

    main_table = Table(main_data, colWidths=col_widths)
    # Merge cells for headers
    table_style = [
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        # Merge header row 1: "Tick which evidence method" spans columns 2-4
        ('SPAN', (2, 0), (4, 0)),
        # Merge "Unit" column header (rows 0-2)
        ('SPAN', (0, 0), (0, 2)),
        # Merge "LO/Assessment criteria" column header (rows 0-2)
        ('SPAN', (1, 0), (1, 2)),
        # Merge "Record of observed activity" spans columns 2-4 in row 2
        ('SPAN', (2, 2), (4, 2)),
        # Merge narrative cells in each data row (columns 2-4)
    ]

    # Merge data row narrative cells (columns 2-4 for each data row starting at index 3)
    for row_idx in range(3, len(main_data)):
        table_style.append(('SPAN', (2, row_idx), (4, row_idx)))
        # Light gray background for header rows
    for row_idx in range(3):
        table_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), colors.lightgrey))
        table_style.append(('FONTNAME', (0, row_idx), (-1, row_idx), 'Helvetica-Bold'))

    # Hide internal horizontal lines between data rows (same as Word doc)
    data_start = 3
    data_end = len(main_data) - 1
    if data_end >= data_start:
        for row_idx in range(data_start, data_end + 1):
            if row_idx > data_start:
                table_style.append(('LINEABOVE', (0, row_idx), (-1, row_idx), 0, colors.white))
            if row_idx < data_end:
                table_style.append(('LINEBELOW', (0, row_idx), (-1, row_idx), 0, colors.white))

    main_table.setStyle(TableStyle(table_style))
    story.append(main_table)

    # --- DECLARATION ---
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "I confirm that the evidence listed above is true record of the performed activities "
        "observed during this assessment.",
        cell_style
    ))

    # --- SIGNATURES GRID (Table 4) ---
    sig_data = [
        [Paragraph("Work base Assessor / witness Signature: __________________", cell_style),
         Paragraph(f"Date: {date}", cell_style)],
        [Paragraph("Candidate's Signature: __________________", cell_style),
         Paragraph(f"Date: {date}", cell_style)],
        [Paragraph("Assessor's Signature: __________________", cell_style),
         Paragraph(f"Date: {date}", cell_style)],
        [Paragraph("Internal Verifier's Signature: __________________", cell_style),
         Paragraph(f"Date: {date}", cell_style)],
    ]
    sig_table = Table(sig_data, colWidths=[4.5 * inch, 2.0 * inch])
    sig_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(sig_table)

    return story


def export_to_pdf(name, date, report_text, assessor_name, timeline="N/A", atmosphere="N/A", selected_pcs=None):
    """
    Generates the official NSQ form as a PDF (matching the Word document layout).
    """
    # Parse selected_pcs for units summary (same logic as Word export)
    units_list = []
    if selected_pcs and isinstance(selected_pcs, list):
        for pc_str in selected_pcs:
            parts = pc_str.split(' - ')
            if len(parts) >= 3:
                units_list.append(parts[0].strip())
    elif isinstance(selected_pcs, str) and selected_pcs:
        units_list = [u.strip() for u in selected_pcs.split(',') if u.strip()]

    units_summary = ", ".join(dict.fromkeys([get_unit_number(u) for u in units_list]))

    # Inject units_summary into the PDF via a marker in report_text
    # The Word template uses units_summary in the metadata table
    # For PDF, we pass it through a special key by temporarily appending to the doc title
    bio = BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=f"NSQ Assessment Report - {name}",
        author=assessor_name
    )

    story = _build_nsq_pdf_story(
        name=name,
        date=date,
        report_text=report_text,
        assessor_name=assessor_name,
        units_summary=units_summary,
        evidence_type="observation"
    )
    doc.build(story)
    return bio.getvalue()


def create_zip_from_reports(reports, table_type):
    """
    Takes a list of fetched report dictionaries (with full text) and the table_type.
    Returns a ZIP file containing the Word documents as bytes.
    """
    bio = BytesIO()
    with zipfile.ZipFile(bio, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for r in reports:
            raw_date = r.get('assessment_date') or r.get('created_at', 'N/A')
            display_date = raw_date.split('T')[0] if 'T' in str(raw_date) else raw_date
            
            assessor_name = (r.get('user_profiles') or {}).get('full_name', 'Unknown Assessor')
            student_name = r.get('student_name') or r.get('candidate_name', 'Student')
            unit_codes = r.get('unit_codes', 'N/A')
            text_content = r.get('text_content', '')
            
            if table_type == "Assessment Reports":
                doc_bytes = export_to_word(
                    name=student_name, 
                    date=display_date, 
                    report_text=text_content, 
                    assessor_name=assessor_name,
                    selected_pcs=unit_codes
                )
            elif table_type == "Personal Statements":
                doc_bytes = export_personal_statement_to_word(
                    name=student_name,
                    date=display_date,
                    statement_text=text_content,
                    selected_pcs=unit_codes
                )
            elif table_type == "Witness Statements":
                witness_name = r.get('witness_name', 'Witness')
                witness_role = r.get('witness_role', 'Supervisor')
                doc_bytes = export_witness_to_word(
                    witness_name=witness_name,
                    witness_role=witness_role,
                    candidate_name=student_name,
                    date=display_date,
                    statement_text=text_content,
                    selected_pcs=unit_codes
                )
            else:
                continue
                
            safe_name = re.sub(r'[^A-Za-z0-9_\-\.]', '_', student_name)
            report_id = r.get('id', '')
            filename = f"NSQ_{safe_name}_{report_id}.docx"
            zipf.writestr(filename, doc_bytes)
            
    return bio.getvalue()
