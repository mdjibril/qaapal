import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import datetime
import sqlite3

# --- CONFIGURATION ---
st.set_page_config(page_title="NSQ Report Architect", layout="wide")

if 'requests_left' not in st.session_state:
    st.session_state.requests_left = 15 

# --- WORD EXPORT FUNCTION ---
def export_to_word(name, date, time, atmosphere, report_text, selected_pcs):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = f"Candidate Name: {name}"
    table.rows[0].cells[1].text = f"Date: {date}"
    table.rows[1].cells[0].text = f"Timeline: {time}"
    table.rows[1].cells[1].text = f"Assessor: Jibril Dauda Muhammad"
    table.rows[2].cells[0].text = f"Atmosphere: {atmosphere}"
    table.rows[2].cells[1].text = "Status: Competent (Progressing)"

    doc.add_paragraph("\n")
    doc.add_heading('Observation Narrative & Evidence', level=1)
    doc.add_paragraph(report_text)
    
    # --- FORMAL CRITERIA SUMMARY SECTION ---
    doc.add_paragraph("\n")
    doc.add_heading('Mapped Performance Criteria Summary', level=2)
    
    units_dict = {}
    for pc in selected_pcs:
        # Split "Unit Code - PC Code: Desc"
        parts = pc.split(' - ')
        unit = parts[0]
        criterion = parts[1].split(':')[0] # Gets just the "PC X.X"
        if unit not in units_dict:
            units_dict[unit] = []
        units_dict[unit].append(criterion)

    for unit, pcs in units_dict.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{unit}: ").bold = True
        p.add_run(", ".join(pcs))

    doc.add_paragraph("\n")
    doc.add_heading('Assessor Signature: _______________________', level=3)
    doc.add_paragraph(f"Verified by Jibril Dauda Muhammad (QAA/2363/ICT) on {date}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- DATABASE UTILITY FUNCTIONS ---
def get_trades():
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM trades")
    data = cursor.fetchall()
    conn.close()
    return data

def get_nested_nos(trade_id):
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, code, title FROM units WHERE trade_id = ?", (trade_id,))
    units = cursor.fetchall()
    nested_data = {}
    for u_id, u_code, u_title in units:
        unit_label = f"{u_code}: {u_title}"
        nested_data[unit_label] = {}
        cursor.execute("SELECT id, lo_num, desc FROM learning_outcomes WHERE unit_id = ?", (u_id,))
        los = cursor.fetchall()
        for lo_id, lo_num, lo_desc in los:
            lo_label = f"{lo_num}: {lo_desc}"
            cursor.execute("SELECT pc_code, desc FROM performance_criteria WHERE lo_id = ?", (lo_id,))
            pcs = [f"{row[0]}: {row[1]}" for row in cursor.fetchall()]
            nested_data[unit_label][lo_label] = pcs
    conn.close()
    return nested_data

def save_report_to_history(name, trade_id, unit_codes, text, date):
    conn = sqlite3.connect('nsq_audit.db')
    conn.execute("""
        INSERT INTO assessment_reports (student_name, trade_id, unit_codes, report_text, assessment_date)
        VALUES (?, ?, ?, ?, ?)
    """, (name, trade_id, unit_codes, text, str(date)))
    conn.commit()
    conn.close()

# --- SIDEBAR ---
# Note: Ensure you have your actual API key here or in the UI
api_key = st.sidebar.text_input("AIzaSyBqxd-ricU_Mp--Un3e2l5GpNRI5-mFg88", type="password")

if api_key:
    # genai.configure(api_key=api_key)
    # model = genai.GenerativeModel('gemini-1.5-flash')
    # st.sidebar.metric("Requests Left", st.session_state.requests_left)
    try:
        genai.configure(api_key=api_key)
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        target_model = 'models/gemini-1.5-flash' 
        if target_model not in available_models:
            flash_models = [m for m in available_models if 'flash' in m]
            target_model = flash_models[0] if flash_models else available_models[0]
        
        model = genai.GenerativeModel(target_model)
        st.sidebar.success(f"Connected: {target_model}")
        st.sidebar.metric("Requests Left (RPM)", st.session_state.requests_left)
    except Exception as e:
        st.sidebar.error(f"Error: {e}")

# --- UI DESIGN ---
st.title("📝 NSQ Report Generator")

trades_list = get_trades()
if trades_list:
    trade_names = [t[1] for t in trades_list]
    selected_trade_name = st.sidebar.selectbox("Select Trade", trade_names)
    selected_trade_id = [t[0] for t in trades_list if t[1] == selected_trade_name][0]
    NOS_DATA = get_nested_nos(selected_trade_id)
else:
    st.stop()

with st.expander("Step 1: Student & Session Info", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Candidate Name")
        assessment_date = st.date_input("Assessment Date", datetime.date.today())
    with col2:
        time_frame = st.text_input("Timeline", placeholder="e.g. 9:00AM – 12:00PM")
        atmosphere = st.text_area("Atmospheric Details", "Moderate temperature, quiet lab environment.")

st.markdown("### Step 2: Select Achieved PCs")
selected_pcs = []
tabs = st.tabs(list(NOS_DATA.keys()))
for i, unit_key in enumerate(NOS_DATA.keys()):
    with tabs[i]:
        for lo_key, pcs in NOS_DATA[unit_key].items():
            with st.expander(lo_key):
                for pc in pcs:
                    if st.checkbox(pc, key=f"{unit_key}_{lo_key}_{pc}"):
                        # Full string stored for prompt accuracy
                        selected_pcs.append(f"{unit_key.split(':')[0]} - {pc}")

learning_moment = st.text_area("Step 3: Learning Moment", placeholder="Key breakthroughs...")

# --- GENERATION LOGIC ---
if st.button("Generate & Finalize Report"):
    if not selected_pcs:
        st.warning("Select at least one PC.")
    else:
        with st.spinner("AI is synthesizing..."):
            # Prepare detailed criteria list for the AI to prevent hallucination
            detailed_criteria_text = "\n".join(selected_pcs)
            
            prompt = f"""
            You are a professional NSQ Assessor. Write a formal narrative assessment report for {student_name} based on the NSQ Performance Evidence Record.

            STRICT DATA (Mention actions for every one of these):
            {detailed_criteria_text}

            CONTEXT:
            - Timeline: {time_frame}
            - Observation: {learning_moment}
            - Environmental Details: {atmosphere}
            
            RULES:
            1. Minimum 10 paragraphs. NO numbering.
            2. Narrative Style: Chronological arrival to completion (e.g., "{student_name} commenced the session at {time_frame}...").
            3. SHORTHAND MAPPING: End every paragraph with grouped codes, e.g. (Unit 8-PC 3.1, 3.2).
            4. NO HALLUCINATION: Do not invent PC numbers outside of the provided list.
            5. TONE: Professional, technical, and objective.
            """
            
            try:
                response = model.generate_content(prompt)
                ai_narrative = response.text
                
                # --- PYTHON-GENERATED TRUTH SUMMARY ---
                summary_block = "\n\n----- SUMMARY OF CRITERIA COVERED -----\n\n"
                u_dict = {}
                for item in selected_pcs:
                    u_code = item.split(' - ')[0]
                    pc_code = item.split(' - ')[1].split(':')[0]
                    if u_code not in u_dict: u_dict[u_code] = []
                    u_dict[u_code].append(pc_code)
                
                for u, pc_list in u_dict.items():
                    summary_block += f"{u}: {', '.join(pc_list)}\n"
                
                # Final combined text
                full_report_text = ai_narrative + summary_block
                
                # Save to History
                units_str = ", ".join(u_dict.keys())
                save_report_to_history(student_name, selected_trade_id, units_str, full_report_text, assessment_date)
                
                st.markdown("### Preview")
                st.info(full_report_text)
                
                # Prepare Word Doc
                doc_bytes = export_to_word(student_name, assessment_date, time_frame, atmosphere, full_report_text, selected_pcs)
                
                c1, c2 = st.columns(2)
                with c1:
                    st.download_button("📥 Download Word (.docx)", data=doc_bytes, file_name=f"NSQ_{student_name.replace(' ', '_')}.docx")
                with c2:
                    st.download_button("Download Text (.txt)", full_report_text, file_name=f"{student_name}.txt")
                st.session_state.requests_left -= 1
                
            except Exception as e:
                st.error(f"Error: {e}")