import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import datetime

# --- CONFIGURATION ---
st.set_page_config(page_title="NSQ Report Architect", layout="wide")

# --- QUOTA TRACKER INITIALIZATION ---
if 'requests_left' not in st.session_state:
    st.session_state.requests_left = 15  # Gemini Free Tier RPM (Requests Per Minute) limit

# --- WORD EXPORT FUNCTION ---
def export_to_word(name, date, time, atmosphere, report_text):
    doc = Document()
    
    # Customizing Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Heading
    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 # Center

    # Student Info Table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = f"Candidate Name: {name}"
    cells[1].text = f"Date: {date}"
    
    cells = table.rows[1].cells
    cells[0].text = f"Timeline: {time}"
    cells[1].text = f"Assessor: Jibril Dauda Muhammad"
    
    cells = table.rows[2].cells
    cells[0].text = f"Atmosphere: {atmosphere}"
    cells[1].text = "Status: Competent (Progressing)"

    doc.add_paragraph("\n") # Spacer
    doc.add_heading('Observation Narrative & Evidence', level=1)
    
    # Narrative Body
    doc.add_paragraph(report_text)
    
    doc.add_paragraph("\n")
    doc.add_heading('Assessor Signature: _______________________', level=3)

    # Save to buffer
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- SIDEBAR: API & QUOTA ---
api_key = st.sidebar.text_input("AIzaSyBqxd-ricU_Mp--Un3e2l5GpNRI5-mFg88", type="password")

if api_key:
    try:
        genai.configure(api_key=api_key)
        # Dynamic Model Fetching
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        target_model = 'models/gemini-1.5-flash' 
        if target_model not in available_models:
            flash_models = [m for m in available_models if 'flash' in m]
            target_model = flash_models[0] if flash_models else available_models[0]
        
        model = genai.GenerativeModel(target_model)
        st.sidebar.success(f"Connected: {target_model}")
        st.sidebar.metric("Requests Left (RPM)", st.session_state.requests_left)
        
    except Exception as e:
        st.sidebar.error(f"Error: {e}")

# --- DATA STRUCTURE ---
import sqlite3

# --- DATABASE UTILITY FUNCTIONS ---
def get_trades():
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM trades")
    data = cursor.fetchall()
    conn.close()
    return data

def get_nos_for_trade(trade_id):
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    # Get all units for the selected trade
    cursor.execute("SELECT id, title FROM units WHERE trade_id = ?", (trade_id,))
    units = cursor.fetchall()
    
    full_nos = {}
    for unit_id, unit_title in units:
        cursor.execute("SELECT pc_code, description FROM performance_criteria WHERE unit_id = ?", (unit_id,))
        pcs = [f"{row[0]}: {row[1]}" for row in cursor.fetchall()]
        full_nos[unit_title] = pcs
    
    conn.close()
    return full_nos

# --- UI DESIGN (Updated for Trade Selection) ---
st.title("📝 NSQ Report Generator")

# 1. Trade Selection Sidebar
trades = get_trades()
trade_names = [t[1] for t in trades]
selected_trade_name = st.sidebar.selectbox("Select Trade", trade_names)
selected_trade_id = [t[0] for t in trades if t[1] == selected_trade_name][0]

# 2. Fetch NOS dynamically
NOS_DATA = get_nos_for_trade(selected_trade_id)

st.subheader(f"Assessor Tool: {selected_trade_name}")

# --- UI DESIGN ---
# st.title("📝 NSQ Report Generator")
# st.subheader("ICT Level 3: Computer Hardware Repairs & Maintenance")

with st.expander("Step 1: Student & Session Info", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Candidate Name", placeholder="e.g. Suleman Dalhatu")
        assessment_date = st.date_input("Assessment Date", datetime.date.today())
    with col2:
        time_frame = st.text_input("Timeline", placeholder="e.g. 9:10AM – 2:10PM")
        atmosphere = st.text_area("Atmospheric Details", "The lab temperature was moderate as the fan was cooling it")

st.markdown("### Step 2: Select Achieved Performance Criteria (PCs)")
selected_pcs = []
tabs = st.tabs(list(NOS_DATA.keys()))
for i, unit in enumerate(NOS_DATA.keys()):
    with tabs[i]:
        for pc in NOS_DATA[unit]:
            if st.checkbox(pc, key=f"{unit}_{pc}"):
                selected_pcs.append(pc)

learning_moment = st.text_area("Step 3: Learning Moment / Hook", placeholder="Describe a specific struggle or breakthrough (e.g., struggling with motherboard screws)...")

# --- GENERATION & EXPORT LOGIC ---
if st.button("Generate & Finalize Report"):
    if not api_key:
        st.error("Please enter your API Key in the sidebar.")
    elif st.session_state.requests_left <= 0:
        st.warning("Rate limit reached. Please wait a minute before generating another report.")
    elif not selected_pcs:
        st.warning("Please select at least one Performance Criterion.")
    else:
        with st.spinner("AI is synthesizing your observation into a professional report..."):
            prompt = f"""
            Write a formal 10-paragraph narrative assessment report for {student_name}.
            
            STRICT CONTEXT:
            - Timeline: {time_frame}
            - Environmental Observation: {atmosphere}
            - Achieved Standards: {', '.join(selected_pcs)}
            - Specific Learning Incident: {learning_moment}
            
            RULES:
            1. Use professional NSQ/NBTE terminology.
            2. End every paragraph with the relevant PC codes in brackets, e.g., (Unit 4: PC 1.1).
            3. Chronological flow: Arrival, Setup, Activity, Challenge, Resolution, Documentation.
            4. Do not include a preamble; start directly with the arrival.
            """
            
            try:
                response = model.generate_content(prompt)
                report_text = response.text
                
                # Update Quota
                st.session_state.requests_left -= 1
                
                st.markdown("---")
                st.markdown("### Report Preview")
                st.info(report_text)
                
                # Create Word Doc
                doc_bytes = export_to_word(student_name, assessment_date, time_frame, atmosphere, report_text)
                
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    st.download_button(
                        label="📥 Download Official Word (.docx)",
                        data=doc_bytes,
                        file_name=f"NSQ_Report_{student_name.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col_dl2:
                    st.download_button("Download Raw Text (.txt)", report_text, file_name=f"{student_name}_Report.txt")
                    
            except Exception as e:
                st.error(f"AI Generation Failed: {e}")