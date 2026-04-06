import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import datetime
import sqlite3

# --- CONFIGURATION ---
st.set_page_config(page_title="NSQ Report Architect", layout="wide")

# --- QUOTA TRACKER INITIALIZATION ---
if 'requests_left' not in st.session_state:
    st.session_state.requests_left = 15 

# --- WORD EXPORT FUNCTION ---
def export_to_word(name, date, time, atmosphere, report_text):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = f"Candidate Name: {name}"
    table.rows[0].cells[1].text = f"Date: {date}"
    table.rows[1].cells[0].text = f"Timeline: {time}"
    table.rows[1].cells[1].text = f"Assessor: Jibril Dauda Muhammad"
    table.rows[2].cells[0].text = f"Atmosphere: {atmosphere}"
    table.rows[2].cells[1].text = "Status: Competent (Progressing)"

    doc.add_paragraph("\n")
    doc.add_heading('Observation Narrative & Evidence', level=1)
    doc.add_paragraph(report_text)
    doc.add_paragraph("\n")
    doc.add_heading('Assessor Signature: _______________________', level=3)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- DATABASE UTILITY FUNCTIONS ---
def get_trades():
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM trades")
    data = cursor.fetchall()
    conn.close()
    return data

def get_nested_nos(trade_id):
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    
    # Get Units
    cursor.execute("SELECT id, code, title FROM units WHERE trade_id = ?", (trade_id,))
    units = cursor.fetchall()
    
    nested_data = {}
    for u_id, u_code, u_title in units:
        unit_label = f"{u_code}: {u_title}"
        nested_data[unit_label] = {}
        
        # Get LOs for this Unit
        cursor.execute("SELECT id, lo_num, desc FROM learning_outcomes WHERE unit_id = ?", (u_id,))
        los = cursor.fetchall()
        for lo_id, lo_num, lo_desc in los:
            lo_label = f"{lo_num}: {lo_desc}"
            
            # Get PCs for this LO
            cursor.execute("SELECT pc_code, desc FROM performance_criteria WHERE lo_id = ?", (lo_id,))
            pcs = [f"{row[0]}: {row[1]}" for row in cursor.fetchall()]
            nested_data[unit_label][lo_label] = pcs
            
    conn.close()
    return nested_data

# --- SIDEBAR: API & QUOTA ---
# Note: I've left the label as 'api_key' for security. 
# Don't forget to paste your key starting with 'AIza...' in the UI.
api_key = st.sidebar.text_input("AIzaSyBqxd-ricU_Mp--Un3e2l5GpNRI5-mFg88", type="password")

if api_key:
    try:
        genai.configure(api_key=api_key)
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        target_model = 'models/gemini-1.5-flash' 
        if target_model not in available_models:
            flash_models = [m for m in available_models if 'flash' in m]
            target_model = flash_models[0] if flash_models else available_models[0]
        
        model = genai.GenerativeModel(target_model)
        st.sidebar.success(f"Connected: {target_model}")
        st.sidebar.metric("Requests Left (RPM)", st.session_state.requests_left)
    except Exception as e:
        st.sidebar.error(f"Error: {e}")

# --- UI DESIGN: TRADE SELECTION ---
st.title("📝 NSQ Report Generator")

trades_list = get_trades()
if trades_list:
    trade_names = [t[1] for t in trades_list]
    selected_trade_name = st.sidebar.selectbox("Select Trade", trade_names)
    selected_trade_id = [t[0] for t in trades_list if t[1] == selected_trade_name][0]

    # Fetch Nested NOS Data
    NOS_DATA = get_nested_nos(selected_trade_id)
    st.subheader(f"Assessor Tool: {selected_trade_name}")
else:
    st.error("No trades found in database. Please run setup_db.py first.")
    st.stop()

# --- UI DESIGN: STUDENT INFO ---
with st.expander("Step 1: Student & Session Info", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Candidate Name", placeholder="e.g. Suleman Dalhatu")
        assessment_date = st.date_input("Assessment Date", datetime.date.today())
    with col2:
        time_frame = st.text_input("Timeline", placeholder="e.g. 9:10AM – 2:10PM")
        atmosphere = st.text_area("Atmospheric Details", "The lab temperature was moderate as the fan was cooling it")

# --- UI DESIGN: NOS CHECKLIST ---
st.markdown("### Step 2: Select Achieved Performance Criteria (PCs)")
selected_pcs = []

if NOS_DATA:
    tabs = st.tabs(list(NOS_DATA.keys()))
    for i, unit_key in enumerate(NOS_DATA.keys()):
        with tabs[i]:
            for lo_key, pcs in NOS_DATA[unit_key].items():
                with st.expander(lo_key):
                    for pc in pcs:
                        # Display individual checkboxes for each PC
                        if st.checkbox(pc, key=f"{unit_key}_{lo_key}_{pc}"):
                            # We store a simplified version for the AI prompt
                            selected_pcs.append(f"{unit_key.split(':')[0]} - {pc}")
else:
    st.warning("No Unit data available for this trade.")

learning_moment = st.text_area("Step 3: Learning Moment / Hook", placeholder="Describe a specific struggle or breakthrough...")

# --- GENERATION & EXPORT LOGIC ---
# if st.button("Generate & Finalize Report"):
#     if not api_key:
#         st.error("Please enter your API Key in the sidebar.")
#     elif st.session_state.requests_left <= 0:
#         st.warning("Rate limit reached. Please wait a minute.")
#     elif not selected_pcs:
#         st.warning("Please select at least one Performance Criterion.")
#     else:
#         with st.spinner("AI is synthesizing your observation into a professional report..."):
#             prompt = f"""
#             Write a formal 10-paragraph narrative assessment report for {student_name}.
            
#             STRICT CONTEXT:
#             - Timeline: {time_frame}
#             - Environmental Observation: {atmosphere}
#             - Achieved Standards: {', '.join(selected_pcs)}
#             - Specific Learning Incident: {learning_moment}
            
#             RULES:
#             1. Use professional NSQ/NBTE terminology.
#             2. End every paragraph with the relevant codes in brackets, e.g., (Unit 4: PC 1.1).
#             3. Chronological flow: Arrival, Setup, Activity, Challenge, Resolution, Documentation.
#             4. Tone: Objective and formal.
#             """
            
#             try:
#                 response = model.generate_content(prompt)
#                 report_text = response.text
#                 st.session_state.requests_left -= 1
                
#                 st.markdown("---")
#                 st.markdown("### Report Preview")
#                 st.info(report_text)
                
#                 doc_bytes = export_to_word(student_name, assessment_date, time_frame, atmosphere, report_text)
                
#                 c1, c2 = st.columns(2)
#                 with c1:
#                     st.download_button("📥 Download Word (.docx)", data=doc_bytes, file_name=f"NSQ_{student_name.replace(' ', '_')}.docx")
#                 with c2:
#                     st.download_button("Download Text (.txt)", report_text, file_name=f"{student_name}.txt")
                    
#             except Exception as e:
#                 st.error(f"AI Generation Failed: {e}")
# ... (Keep Imports and Word Export functions from previous version) ...

# --- GENERATION LOGIC (Inside Home.py) ---
if st.button("Generate & Finalize Report"):
    if not api_key:
        st.error("Please enter your API Key in the sidebar.")
    elif st.session_state.requests_left <= 0:
        st.warning("Rate limit reached. Please wait a minute.")
    elif not selected_pcs:
        st.warning("Please select at least one Performance Criterion.")
    else:
        with st.spinner("AI is synthesizing your observation..."):
            # REFINED PROMPT FOR CLEANER OUTPUT
            # prompt = f"""
            # You are a professional NSQ Assessor. Write a formal narrative assessment report for {student_name}.
            
            # CONTEXT:
            # - Timeline: {time_frame}
            # - Environment: {atmosphere}
            # - Achieved PCs: {', '.join(selected_pcs)}
            # - Observation: {learning_moment}
            
            # STRICT FORMATTING RULES:
            # 1. Length: Minimum 10 paragraphs. Do not exceed 15 unless the observation is highly complex.
            # 2. NO NUMBERING: Do not number the paragraphs. Use clean, professional prose.
            # 3. SHORTHAND MAPPING: At the end of every paragraph, list the relevant PCs in this EXACT format: (Unit X-PC Y.Y, Z.Z). 
            #    Example: If PCs are from Unit 8, write (Unit 8-PC 3.1, 3.2, 3.4).
            # 4. TONE: Objective and technical. Use terms like "ESD compliance," "POST diagnostic," and "thermal management."
            # 5. FLOW: Chronological (Arrival -> Setup -> Execution -> Troubleshooting -> Final Check).
            # """
            # Updated Prompt Logic in Home.py
            prompt = f"""
            Write a formal narrative assessment report for {student_name} based on the NSQ Performance Evidence Record.

            CONTEXT:
            - Units: {', '.join(set([pc.split('-')[0] for pc in selected_pcs]))}
            - Observation: {learning_moment}
            - Environmental Details: {atmosphere}

            RULES:
            1. Minimum 10 paragraphs. NO paragraph numbering.
            2. Narrative Style: Start directly with the arrival (e.g., "{student_name} commenced the session at {time_frame}...").
            3. SHORTHAND MAPPING: At the end of every paragraph, group the relevant PCs by unit.
            FORMAT: (Unit [Number]-PC [Code], [Code]) 
            EXAMPLE: (Unit 8-PC 3.1, 3.2, 5.1)
            4. TONE: Professional, technical, and objective.
            """
            
            try:
                response = model.generate_content(prompt)
                report_text = response.text
                
                st.session_state.requests_left -= 1
                st.markdown("---")
                st.markdown("### Report Preview")
                st.info(report_text)
                
                doc_bytes = export_to_word(student_name, assessment_date, time_frame, atmosphere, report_text)
                
                c1, c2 = st.columns(2)
                with c1:
                    st.download_button("📥 Download Word (.docx)", data=doc_bytes, file_name=f"NSQ_{student_name.replace(' ', '_')}.docx")
                with c2:
                    st.download_button("Download Text (.txt)", report_text, file_name=f"{student_name}.txt")
                    
            except Exception as e:
                st.error(f"AI Generation Failed: {e}")
# ... (Keep Imports and Word Export functions from previous version) ...