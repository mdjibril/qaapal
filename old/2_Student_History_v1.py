# import streamlit as st
# import sqlite3

# st.set_page_config(page_title="Student History", layout="wide")
# st.title("📜 Assessment History")

# conn = sqlite3.connect('nsq_audit.db')
# search = st.text_input("Search Student Name")

# # Query the new assessment_reports table
# query = "SELECT student_name, unit_codes, assessment_date, report_text FROM assessment_reports"
# if search:
#     query += f" WHERE student_name LIKE '%{search}%'"
# query += " ORDER BY timestamp DESC"

# reports = conn.execute(query).fetchall()

# for r in reports:
#     with st.expander(f"{r[2]} - {r[0]} (Units: {r[1]})"):
#         st.write(r[3]) # This displays the full report + the summary
        
# conn.close()

# -------------------

# import streamlit as st
# import sqlite3

# st.set_page_config(page_title="Student History", layout="wide")

# st.title("📜 Assessment History")

# conn = sqlite3.connect('nsq_audit.db')
# search_q = st.text_input("Search by Student Name", placeholder="e.g. Hauwa Adamu")

# query = "SELECT * FROM assessment_reports"
# if search_q:
#     query += f" WHERE student_name LIKE '%{search_q}%'"
# query += " ORDER BY timestamp DESC"

# reports = conn.execute(query).fetchall()

# if not reports:
#     st.info("No reports found.")
# else:
#     for r in reports:
#         with st.expander(f"{r[5]} - {r[1]} (Units: {r[3]})"):
#             st.caption(f"Trade ID: {r[2]} | Generated on: {r[6]}")
#             st.write(r[4])
            
#             # Option to re-download
#             if st.button("Prepare Download", key=f"dl_{r[0]}"):
#                 # You can call your export_to_word function here
#                 st.write("Feature: Re-exporting coming soon...")

# conn.close()

# --------------------

import streamlit as st
import sqlite3
from docx import Document
from docx.shared import Pt
from io import BytesIO

st.set_page_config(page_title="Student History", layout="wide")

# --- REUSE THE EXPORT FUNCTION ---
def export_to_word(name, date, unit_codes, report_text):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 

    # History Table (Simplified based on stored data)
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = f"Candidate Name: {name}"
    table.rows[0].cells[1].text = f"Date: {date}"
    table.rows[1].cells[0].text = f"Units Covered: {unit_codes}"
    table.rows[1].cells[1].text = f"Assessor: Jibril Dauda Muhammad"

    doc.add_paragraph("\n")
    doc.add_heading('Observation Narrative & Evidence', level=1)
    doc.add_paragraph(report_text)
    
    doc.add_paragraph("\n")
    doc.add_heading('Assessor Signature: _______________________', level=3)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- MAIN PAGE LOGIC ---
st.title("📜 Assessment History")
st.info("Search and re-download previously generated NSQ reports.")

conn = sqlite3.connect('nsq_audit.db')
search_q = st.text_input("Search by Student Name", placeholder="e.g. Hauwa Adamu")

# Table structure: 0:id, 1:name, 2:trade_id, 3:units, 4:text, 5:date, 6:timestamp
query = "SELECT * FROM assessment_reports"
if search_q:
    query += f" WHERE student_name LIKE '%{search_q}%'"
query += " ORDER BY timestamp DESC"

reports = conn.execute(query).fetchall()

if not reports:
    st.info("No reports found in the database.")
else:
    for r in reports:
        # r[5] is assessment_date, r[1] is student_name, r[3] is unit_codes
        with st.expander(f"📅 {r[5]} | 👤 {r[1]} (Units: {r[3]})"):
            st.caption(f"Database ID: {r[0]} | Record Created: {r[6]}")
            
            # st.markdown("""
            #     <style>
            #     div[data-testid="stTextarea"] .st-key-my_report [data-baseweb="base-input"] [disabled] {
            #         background-color: #fff4e6 !important;
            #         -webkit-text-fill-color: #854d0e !important;
            #     }
            #     </style>
            # """, unsafe_allow_html=True)

            # Show a preview of the text
            st.text_area("Report Preview", value=r[4], height=200, disabled=True)
            
            # --- DOWNLOAD LOGIC ---
            # Generate the doc bytes specifically for this record
            doc_bytes = export_to_word(
                name=r[1], 
                date=r[5], 
                unit_codes=r[3], 
                report_text=r[4]
            )
            
            st.download_button(
                label="📥 Download Word Document",
                data=doc_bytes,
                file_name=f"NSQ_Report_{r[1].replace(' ', '_')}_{r[5]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{r[0]}" # Unique key per report
            )

conn.close()