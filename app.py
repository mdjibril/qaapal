import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import datetime
import sqlite3

# --- CONFIGURATION ---
st.set_page_config(page_title="NSQ Report Architect", layout="wide")

if 'requests_left' not in st.session_state:
    st.session_state.requests_left = 15 

# --- WORD EXPORT FUNCTION ---
def export_to_word(name, date, time, atmosphere, report_text, selected_pcs, assessor_name, assessor_id):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = f"Candidate Name: {name}"
    table.rows[0].cells[1].text = f"Date: {date}"
    table.rows[1].cells[0].text = f"Timeline: {time}"
    table.rows[1].cells[1].text = f"Assessor: {assessor_name}"
    table.rows[2].cells[0].text = f"Atmosphere: {atmosphere}"
    table.rows[2].cells[1].text = "Status: Competent (Progressing)"

    doc.add_paragraph("\n")
    doc.add_heading('Observation Narrative & Evidence', level=1)
    doc.add_paragraph(report_text)
    
    # --- FORMAL CRITERIA SUMMARY SECTION ---
    doc.add_paragraph("\n")
    doc.add_heading('Mapped Performance Criteria Summary', level=2)
    
    units_dict = {}
    for pc in selected_pcs:
        parts = pc.split(' - ')
        unit = parts[0]
        criterion = parts[1].split(':')[0] 
        if unit not in units_dict:
            units_dict[unit] = []
        units_dict[unit].append(criterion)

    for unit, pcs in units_dict.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{unit}: ").bold = True
        p.add_run(", ".join(pcs))

    doc.add_paragraph("\n")
    doc.add_heading(f'Assessor Signature: _______________________', level=3)
    doc.add_paragraph(f"Verified by {assessor_name} ({assessor_id}) on {date}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- DATABASE UTILITY FUNCTIONS ---
def get_trades():
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM trades")
    data = cursor.fetchall()
    conn.close()
    return data

def get_nested_nos(trade_id):
    conn = sqlite3.connect('nsq_audit.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, code, title FROM units WHERE trade_id = ?", (trade_id,))
    units = cursor.fetchall()
    nested_data = {}
    for u_id, u_code, u_title in units:
        unit_label = f"{u_code}: {u_title}"
        nested_data[unit_label] = {}
        cursor.execute("SELECT id, lo_num, desc FROM learning_outcomes WHERE unit_id = ?", (u_id,))
        los = cursor.fetchall()
        for lo_id, lo_num, lo_desc in los:
            lo_label = f"{lo_num}: {lo_desc}"
            cursor.execute("SELECT pc_code, desc FROM performance_criteria WHERE lo_id = ?", (lo_id,))
            pcs = [f"{row[0]}: {row[1]}" for row in cursor.fetchall()]
            nested_data[unit_label][lo_label] = pcs
    conn.close()
    return nested_data

def save_report_to_history(name, trade_id, unit_codes, text, date):
    conn = sqlite3.connect('nsq_audit.db')
    conn.execute("""
        INSERT INTO assessment_reports (student_name, trade_id, unit_codes, report_text, assessment_date)
        VALUES (?, ?, ?, ?, ?)
    """, (name, trade_id, unit_codes, text, str(date)))
    conn.commit()
    conn.close()

# --- SIDEBAR ---
st.sidebar.title("Configuration")
api_key = st.sidebar.text_input("AIzaSyBqxd-ricU_Mp--Un3e2l5GpNRI5-mFg88", type="password")

if api_key:
    try:
        genai.configure(api_key=api_key)
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        target_model = 'models/gemini-1.5-flash' 
        if target_model not in available_models:
            flash_models = [m for m in available_models if 'flash' in m]
            target_model = flash_models[0] if flash_models else available_models[0]
        
        model = genai.GenerativeModel(target_model)
        st.sidebar.success(f"Connected: {target_model}")
        # st.sidebar.metric("Requests Left (RPM)", st.session_state.requests_left)
    except Exception as e:
        st.sidebar.error(f"Error: {e}")

# NEW: Dynamic Assessor Details
st.sidebar.markdown("---")
st.sidebar.subheader("Assessor Details")
assessor_name = st.sidebar.text_input("Assessor Name", value="Jibril Dauda Muhammad")
assessor_id = st.sidebar.text_input("Assessor ID", value="QAA/2363/ICT")

# --- SIDEBAR: ENVIRONMENT PRESETS ---
# st.sidebar.markdown("---")
st.sidebar.subheader("Today's Environment")

env_options = {
    "Morning (Cool)": "The morning air was cool and the lab was quiet, providing a focused atmosphere with plenty of natural light.",
    "Afternoon (Warm)": "The lab temperature was moderate; the ceiling fans were active to maintain a comfortable working environment during the peak afternoon heat.",
    "Technical/Busy": "The lab was active with the hum of server fans and multiple workstations in use, creating a realistic, high-energy technical environment.",
    "Rainy/Overcast": "Due to the weather, the lab was lit with overhead fluorescent lights; the atmosphere was cool and calm.",
    "Custom": "" 
}

selected_env_preset = st.sidebar.selectbox("Choose a Preset", list(env_options.keys()))
default_env_text = env_options[selected_env_preset]

# --- UI DESIGN ---
st.title("📝 NSQ Report Generator")

trades_list = get_trades()
if trades_list:
    trade_names = [t[1] for t in trades_list]
    selected_trade_name = st.sidebar.selectbox("Select Trade", trade_names)
    selected_trade_id = [t[0] for t in trades_list if t[1] == selected_trade_name][0]
    NOS_DATA = get_nested_nos(selected_trade_id)
else:
    st.stop()

with st.expander("Step 1: Student & Session Info", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Candidate Name")
        assessment_date = st.date_input("Assessment Date", datetime.date.today())
    with col2:
        time_frame = st.text_input("Timeline", placeholder="e.g. 9:00AM – 12:00PM")
        atmosphere = st.text_area("Atmospheric Details", 
            value=default_env_text, 
            help="You can edit the preset text here if needed for this specific student."
        )

st.markdown("#### Step 2: Select Achieved PCs")
selected_pcs = []
tabs = st.tabs(list(NOS_DATA.keys()))
for i, unit_key in enumerate(NOS_DATA.keys()):
    with tabs[i]:
        for lo_key, pcs in NOS_DATA[unit_key].items():
            with st.expander(lo_key):
                for pc in pcs:
                    if st.checkbox(pc, key=f"{unit_key}_{lo_key}_{pc}"):
                        selected_pcs.append(f"{unit_key.split(':')[0]} - {pc}")

# st.markdown("---")
st.markdown("#### Step 3: Unique Learning Moment / Hook")

# The Formula Guide (A visual reminder for the user)
st.info("""
**💡 Pro-Tip for Unique Reports:** Use the 'Observation Formula' for better AI results:
**[Action]** + **[Specific Tool/Component]** + **[Specific Result or Quote]**
*Example: 'Struggled with the RJ45 crimping tool at first but corrected the pin alignment manually after a second attempt.'*
""")

learning_moment = st.text_area(
    "Observation Notes", 
    placeholder="What did this specific student do, say, or struggle with during this session?",
    height=150,
    help="This is the 'flavor' that makes this report different from the others. Even 1-2 sentences here will make the AI output much more realistic."
)

# learning_moment = st.text_area("Step 3: Learning Moment", placeholder="What did they struggle with? What did they say? What specific tool did they use? (e.g., 'Suleman struggled with the crimping tool at first but mastered it by the third try.')")

# --- GENERATION LOGIC ---
if st.button("Generate & Finalize Report"):
    if not selected_pcs:
        st.warning("Select at least one PC.")
    elif not assessor_name:
        st.warning("Please enter the Assessor's name in the sidebar.")
    else:
        with st.spinner("AI is synthesizing your observation..."):
            # --- FIX: GET FULL UNIT TITLES ---
            # This looks at the checkbox labels and finds the unique Units selected
            # It keeps the whole string before the " - PC" part
            unique_units = []
            for pc in selected_pcs:
                # pc format is "UnitCode: Title - PC Code: Desc"
                unit_part = pc.split(' - ')[0] 
                if unit_part not in unique_units:
                    unique_units.append(unit_part)
            
            unit_header_info = "\n".join(unique_units)

            # FIX 2: Create a very detailed list of PC descriptions for the AI to read
            detailed_criteria_text = "\n".join(selected_pcs)
            
            # FIX 3: Force the date into the prompt so it stops using [Insert Date]
            formatted_date = assessment_date.strftime("%B %d, %Y")

            # prompt = f"""
            # You are a professional NSQ Assessor. Write a formal narrative assessment report.
            
            # REPORT HEADER DATA:
            # - Assessor: {assessor_name} ({assessor_id})
            # - Candidate: {student_name}
            # - Date of Assessment: {formatted_date}
            # - Units of Competence: {unit_header_info}
            # - Timeline: {time_frame}
            # - Environment: {atmosphere}

            # STRICT EVIDENCE TO COVER (Detailed Descriptions):
            # {detailed_criteria_text}

            # SPECIFIC INSTRUCTIONS:
            # 1. START the report immediately with the Header Data provided above.
            # 2. DO NOT use placeholders like "[Insert Date]". Use {formatted_date}.
            # 3. LENGTH: Minimum 10 paragraphs. Expand deeply on the "How" and "Why" for each PC description provided above.
            # 4. NARRATIVE: Describe {student_name}'s specific actions and explanations during the {time_frame} window.
            # 5. SHORTHAND MAPPING: End every paragraph with grouped codes, e.g. (Unit 11-PC 5.1).
            # 6. HOOK: Incorporate this specific observation: {learning_moment}
            # 7. TONE: Professional, technical, and exhaustive.
            # """
            # --- REFINED PROMPT LOGIC ---
            prompt = f"""
            You are a Lead NSQ Assessor. Write a high-level technical narrative for {student_name}.

            REPORT CONTEXT:
            - Assessor: {assessor_name} ({assessor_id})
            - Date: {formatted_date}
            - Units: {unit_header_info}
            - Environment: {atmosphere}
            - The "Hook" (Breaktrough moment): {learning_moment}

            CRITERIA TO INTEGRATE (The Ingredients):
            {detailed_criteria_text}

            STRICT NARRATIVE RULES:
            1. NO STAGED SEQUENCING: Do not write one paragraph per PC. A single paragraph should naturally weave together 2 or 3 related or non-related PCs.
            2. CONTINUOUS FLOW: The report must read like a professional "day-in-the-life" observation. Use phrases like "Simultaneously," "This led the candidate to," or "Building on the previous step.", {student_name} commenced the session at {time_frame}, {student_name} ended the session at {time_frame}.
            3. MINIMUM 8 PARAGRAPHS: Each paragraph must be dense and technical.
            4. MAXIMUM 10 PARAGRAPHS: Do not exceed 10 paragraphs. Be concise but exhaustive.
            5. MAXIMUM PAGES: The final report should not exceed 2 pages when formatted in Word with standard margins and font size.
            6. MAXIMIZE CRITERIA INTEGRATION: Each paragraph should integrate multiple PCs, showing how the candidate's actions and explanations naturally covered several criteria at once.
            7. MAPPING: Use grouped shorthand at the end of paragraphs. 
            Example: (Unit 11-PC 5.1, 5.2, 5.4).
            8. THE HOOK: Do not just "mention" the {learning_moment}; make it a central part of the story where multiple criteria were met during that specific breakthrough.
            9. AVOID PLACEHOLDERS: Use the provided date and names.
            10. TONE: The narrative should be formal, technical, and exhaustive, suitable for a professional assessment report.
            11. DO NOT REPEAT CRITERIA: Each PC should only be referenced once in the narrative, but can be grouped with related PCs for a more natural flow.
            12. USE ALL PROVIDED CRITERIA: Ensure that every PC in the detailed list is integrated into the narrative, either directly or through inference based on the candidate's actions and explanations.
            13. NO HALLUCINATION: Do not invent PC numbers outside of the provided list.
            14. SHOW, DON'T TELL: Instead of saying "the candidate was competent," describe what the candidate did that demonstrated competence, and link those actions to the PCs.
            15. DONT USE HEAVY DICTIONARY LANGUAGE: The report should be professional but also clear and straightforward. Avoid overly complex sentences that could obscure the candidate's actual performance and the criteria met.
            """

            try:
                response = model.generate_content(prompt)
                ai_narrative = response.text
                
                # --- PYTHON-GENERATED TRUTH SUMMARY ---
                summary_block = "\n\n----- SUMMARY OF CRITERIA COVERED -----\n\n"
                u_dict = {}
                for item in selected_pcs:
                    u_code = item.split(' - ')[0].split(':')[0] # Gets just 'ICT/CMR/011/L3'
                    pc_code = item.split(' - ')[1].split(':')[0] # Gets just 'PC 5.1'
                    if u_code not in u_dict: u_dict[u_code] = []
                    u_dict[u_code].append(pc_code)
                
                for u, pc_list in u_dict.items():
                    summary_block += f"{u}: {', '.join(pc_list)}\n"
                
                # Combine everything
                full_report_text = ai_narrative + summary_block
                
                # Save to History
                units_str = ", ".join(u_dict.keys())
                save_report_to_history(student_name, selected_trade_id, units_str, full_report_text, assessment_date)
                
                st.markdown("### Preview")
                st.info(full_report_text)
                
                # Prepare Word Doc
                doc_bytes = export_to_word(
                    student_name, 
                    assessment_date, 
                    time_frame, 
                    atmosphere, 
                    full_report_text, 
                    selected_pcs,
                    assessor_name,
                    assessor_id
                    # name=student_name, 
                    # date=formatted_date, # Use the string date
                    # time=time_frame, 
                    # atmosphere=atmosphere, 
                    # report_text=full_report_text, 
                    # selected_pcs=selected_pcs,
                    # assessor_name=assessor_name,
                    # assessor_id=assessor_id
                )
                
                # Download Buttons
                c1, c2 = st.columns(2)
                with c1:
                    st.download_button("📥 Download Word (.docx)", data=doc_bytes, file_name=f"NSQ_{student_name.replace(' ', '_')}.docx")
                with c2:
                    st.download_button("Download Text (.txt)", full_report_text, file_name=f"{student_name}.txt")
                
            except Exception as e:
                st.error(f"Error: {e}")
                st.session_state.requests_left -= 1