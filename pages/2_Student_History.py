import streamlit as st
import sqlite3
from docx import Document
from docx.shared import Pt
from io import BytesIO

st.set_page_config(page_title="Student History", layout="wide")

# --- REUSE THE EXPORT FUNCTION (Keep this for downloads) ---
def export_to_word(name, date, unit_codes, report_text, assessor_name, assessor_id):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = f"Candidate Name: {name}"
    table.rows[0].cells[1].text = f"Date: {date}"
    table.rows[1].cells[0].text = f"Units Covered: {unit_codes}"
    table.rows[1].cells[1].text = f"Assessor: {assessor_name} ({assessor_id})"
    doc.add_paragraph("\n")
    doc.add_paragraph(report_text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- MAIN PAGE LOGIC ---
st.title("📜 Assessment History & Management")

conn = sqlite3.connect('nsq_audit.db')
search_q = st.text_input("🔍 Search by Student Name", placeholder="e.g. Aisha Kabir")

# Fetch reports
query = "SELECT * FROM assessment_reports"
if search_q:
    query += f" WHERE student_name LIKE '%{search_q}%'"
query += " ORDER BY timestamp DESC"
reports = conn.execute(query).fetchall()

if not reports:
    st.info("No reports found.")
else:
    st.subheader("Manage Records")
    st.caption("Select the records you wish to delete, then scroll to the bottom to confirm.")

    # We use a dictionary to track which IDs are selected for deletion
    delete_dict = {}

    for r in reports:
        # Table structure: 0:id, 1:name, 2:trade_id, 3:units, 4:text, 5:date, 6:timestamp
        col_check, col_content = st.columns([0.1, 0.9])
        
        with col_check:
            # Create a checkbox for each record
            delete_dict[r[0]] = st.checkbox("", key=f"select_{r[0]}")
        
        with col_content:
            with st.expander(f"📅 {r[5]} | 👤 {r[1]} (Units: {r[3]})"):
                st.write(r[4])
                
                # Download Button for this specific record
                doc_bytes = export_to_word(r[1], r[5], r[3], r[4], "Assessor", "ID") # Placeholder names
                st.download_button(
                    label="📥 Download Word",
                    data=doc_bytes,
                    file_name=f"NSQ_{r[1]}_{r[5]}.docx",
                    key=f"dl_{r[0]}"
                )

    st.markdown("---")
    
    # --- DELETION ACTIONS ---
    selected_ids = [report_id for report_id, checked in delete_dict.items() if checked]
    
    if selected_ids:
        st.error(f"You have selected {len(selected_ids)} record(s) for deletion.")
        confirm_del = st.checkbox("I confirm that I want to PERMANENTLY delete these records.")
        
        if st.button("🔥 Delete Selected Records", type="primary", disabled=not confirm_del):
            try:
                # Construct the SQL delete command
                id_placeholders = ', '.join(['?'] * len(selected_ids))
                conn.execute(f"DELETE FROM assessment_reports WHERE id IN ({id_placeholders})", selected_ids)
                conn.commit()
                st.success(f"Successfully deleted {len(selected_ids)} records.")
                st.rerun()
            except Exception as e:
                st.error(f"Error during deletion: {e}")

conn.close()