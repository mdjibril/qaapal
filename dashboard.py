import streamlit as st
import google.generativeai as genai
from groq import Groq 
import requests # Added for OpenRouter
import json
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import datetime
import sqlite3
import time
import google.api_core.exceptions
import database as db  # Import our new module

# --- CONFIGURATION ---
st.set_page_config(page_title="NSQ Report Architect", layout="wide")

if 'last_request_time' not in st.session_state:
    st.session_state.last_request_time = 0

# --- MODULAR AI ROUTER WITH AUTO-DISCOVERY ---
def validate_and_generate(provider, model_name, api_key, prompt=None):
    """
    Handles API calls with dynamic model discovery for Gemini
    to avoid 404 and naming convention errors.
    """
    api_key = api_key.strip()
    
    try:
        if provider == "Gemini":
            genai.configure(api_key=api_key)
            
            # --- START AUTO-DISCOVERY LOGIC ---
            # This replicates your old code's success
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            
            # Try to use what the user selected, but if not found, find the best match
            # e.g., if user selected 'gemini-1.5-flash', look for 'models/gemini-1.5-flash'
            actual_model_name = None
            if any(model_name in m for m in available_models):
                actual_model_name = [m for m in available_models if model_name in m][0]
            else:
                # Fallback to the first available model if selection fails
                actual_model_name = available_models[0]
            # --- END AUTO-DISCOVERY LOGIC ---

            model = genai.GenerativeModel(actual_model_name)
            
            if prompt:
                response = model.generate_content(prompt)
                return response.text
            else:
                return f"✅ Connected: {actual_model_name}"

        elif provider == "Groq":
            client = Groq(api_key=api_key)
            if prompt:
                completion = client.chat.completions.create(
                    model=model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7
                )
                return completion.choices[0].message.content
            else:
                client.models.list()
                return f"✅ Connected: {model_name}"

        elif provider == "OpenRouter":
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost:8501", 
            }
            data = {
                "model": model_name,
                "messages": [{"role": "user", "content": prompt if prompt else "Hello"}]
            }
            response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, data=json.dumps(data))
            if response.status_code == 200:
                return response.json()['choices'][0]['message']['content'] if prompt else f"✅ Connected: {model_name}"
            else:
                return f"API_ERROR: {response.status_code} - {response.text}"

    except Exception as e:
        return f"API_ERROR: {str(e)}"

# --- WORD EXPORT FUNCTION ---
def export_to_word(name, date, time_val, atmosphere, report_text, selected_pcs, assessor_name, assessor_id):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    header = doc.add_heading('National Skills Qualification (NSQ) - Assessment Report', 0)
    header.alignment = 1 
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = f"Candidate Name: {name}"
    table.rows[0].cells[1].text = f"Date: {date}"
    table.rows[1].cells[0].text = f"Timeline: {time_val}"
    table.rows[1].cells[1].text = f"Assessor: {assessor_name}"
    table.rows[2].cells[0].text = f"Atmosphere: {atmosphere}"
    table.rows[2].cells[1].text = "Status: Competent (Progressing)"
    doc.add_paragraph("\n")
    doc.add_heading('Observation Narrative & Evidence', level=1)
    doc.add_paragraph(report_text)
    doc.add_paragraph("\n")
    doc.add_heading('Mapped Performance Criteria Summary', level=2)
    units_dict = {}
    for pc in selected_pcs:
        parts = pc.split(' - ')
        unit = parts[0]
        criterion = parts[1].split(':')[0] 
        if unit not in units_dict: units_dict[unit] = []
        units_dict[unit].append(criterion)
    for unit, pcs in units_dict.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{unit}: ").bold = True
        p.add_run(", ".join(pcs))
    doc.add_paragraph("\n")
    doc.add_heading(f'Assessor Signature: _______________________', level=3)
    doc.add_paragraph(f"Verified by {assessor_name} ({assessor_id}) on {date}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- DB HELPERS ---
# def get_trades():
#     conn = sqlite3.connect('nsq_audit.db')
#     cursor = conn.cursor()
#     cursor.execute("SELECT id, name FROM trades")
#     data = cursor.fetchall()
#     conn.close()
#     return data

# def get_nested_nos(trade_id):
#     conn = sqlite3.connect('nsq_audit.db')
#     cursor = conn.cursor()
#     cursor.execute("SELECT id, code, title FROM units WHERE trade_id = ?", (trade_id,))
#     units = cursor.fetchall()
#     nested_data = {}
#     for u_id, u_code, u_title in units:
#         unit_label = f"{u_code}: {u_title}"
#         nested_data[unit_label] = {}
#         cursor.execute("SELECT id, lo_num, desc FROM learning_outcomes WHERE unit_id = ?", (u_id,))
#         los = cursor.fetchall()
#         for lo_id, lo_num, lo_desc in los:
#             lo_label = f"{lo_num}: {lo_desc}"
#             cursor.execute("SELECT pc_code, desc FROM performance_criteria WHERE lo_id = ?", (lo_id,))
#             pcs = [f"{row[0]}: {row[1]}" for row in cursor.fetchall()]
#             nested_data[unit_label][lo_label] = pcs
#     conn.close()
#     return nested_data

# def save_report_to_history(name, trade_id, unit_codes, text, date):
#     conn = sqlite3.connect('nsq_audit.db')
#     conn.execute("INSERT INTO assessment_reports (student_name, trade_id, unit_codes, report_text, assessment_date) VALUES (?, ?, ?, ?, ?)", (name, trade_id, unit_codes, text, str(date)))
#     conn.commit()
#     conn.close()

# --- SIDEBAR: PROVIDER & CONNECTION CHECK ---
st.sidebar.subheader("📡 AI Provider Settings")
ai_provider = st.sidebar.selectbox("Select Provider", ["Gemini", "Groq", "OpenRouter"])

if ai_provider == "Gemini":
    target_key = st.sidebar.text_input("Gemini API Key", type="password")
    # We list simple names; the Router will find the 'models/xxx' version
    target_model = st.sidebar.selectbox("Gemini Preference", ["gemini-2.5-flash", "gemini-1.5-pro", "gemini-pro"])
elif ai_provider == "Groq":
    target_key = st.sidebar.text_input("Groq API Key", type="password")
    target_model = st.sidebar.selectbox("Groq Model", ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"])
elif ai_provider == "OpenRouter":
    target_key = st.sidebar.text_input("OpenRouter Key", type="password")
    target_model = st.sidebar.selectbox("Model", ["anthropic/claude-3-haiku", "nvidia/nemotron-3-super-120b-a12b:free", "arcee-ai/trinity-large-preview:free", "google/gemma-4-31b-it:free"])

# THE VERIFICATION BUTTON
if target_key:
    if st.sidebar.button("Verify Connection"):
        with st.sidebar:
            with st.spinner("Checking..."):
                res = validate_and_generate(ai_provider, target_model, target_key)
                if "API_ERROR" in str(res):
                    st.error(f"❌ {res}")
                else:
                    st.success(res) # Shows the "✅ Connected: models/..." message

dev_mode = st.sidebar.checkbox("🛠️ Dev Mode (Skip AI)", value=False)

st.sidebar.markdown("---")
st.sidebar.subheader("General Information")
assessor_name = st.sidebar.text_input("Assessor Name", value="Jibril Dauda Muhammad")
assessor_id = st.sidebar.text_input("Assessor ID", value="QAA/XXXX/ICT")

env_options = {
    "Morning (Cool)": "The morning air was cool and the lab was quiet, providing a focused atmosphere with plenty of natural light.",
    "Afternoon (Warm)": "The lab temperature was moderate; the ceiling fans were active to maintain a comfortable working environment during the peak afternoon heat.",
    "Technical/Busy": "The lab was active with the hum of server fans and multiple workstations in use, creating a realistic, high-energy technical environment.",
    "Rainy/Overcast": "Due to the weather, the lab was lit with overhead fluorescent lights; the atmosphere was cool and calm.",
    "Custom": "" 
}
selected_env_preset = st.sidebar.selectbox("Choose a Preset", list(env_options.keys()))
default_env_text = env_options[selected_env_preset]

# --- UI DESIGN ---
st.title("📝 NSQ Report Generator")
st.markdown("#### Step 1: Student & Session Info")
# trades_list = get_trades()
# if trades_list:
#     trade_names = [t[1] for t in trades_list]
#     selected_trade_name = st.sidebar.selectbox("Select Trade", trade_names)
#     selected_trade_id = [t[0] for t in trades_list if t[1] == selected_trade_name][0]
#     NOS_DATA = get_nested_nos(selected_trade_id)
# else:
#     st.stop()

trades = db.fetch_trades()
if not trades.empty:
    selected_trade_name = st.sidebar.selectbox("Select Trade", trades['name'])
    selected_trade_id = trades.loc[trades['name'] == selected_trade_name, 'id'].iloc[0]
    NOS_DATA = db.fetch_nested_nos(selected_trade_id)
    if NOS_DATA:
        st.markdown("#### Step 2: Select Achieved PCs")
        selected_pcs = []
        tabs = st.tabs(list(NOS_DATA.keys()))
        for i, unit_key in enumerate(NOS_DATA.keys()):
            # ... your existing tab logic ...
            with tabs[i]:
                for lo_key, pcs in NOS_DATA[unit_key].items():
                    with st.expander(lo_key):
                        for pc in pcs:
                            if st.checkbox(pc, key=f"{unit_key}_{lo_key}_{pc}"):
                                selected_pcs.append(f"{unit_key.split(':')[0]} - {pc}")
    else:
        st.warning("No units or performance criteria found for this trade.")
else:
    st.error("No trades found in database.")
    st.stop()

with st.expander("Step 1: Details", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Candidate Name")
        assessment_date = st.date_input("Assessment Date", datetime.date.today())
    with col2:
        time_frame = st.text_input("Timeline", placeholder="e.g. 9:00AM – 12:00PM")
        atmosphere = st.text_area("Atmospheric Details", value=default_env_text)

# st.markdown("#### Step 2: Select Achieved PCs")
# selected_pcs = []
# tabs = st.tabs(list(NOS_DATA.keys()))
# for i, unit_key in enumerate(NOS_DATA.keys()):
#     with tabs[i]:
#         for lo_key, pcs in NOS_DATA[unit_key].items():
#             with st.expander(lo_key):
#                 for pc in pcs:
#                     if st.checkbox(pc, key=f"{unit_key}_{lo_key}_{pc}"):
#                         selected_pcs.append(f"{unit_key.split(':')[0]} - {pc}")

st.markdown("#### Step 3: Unique Learning Moment")
# The Formula Guide (A visual reminder for the user)
st.info("""
**💡 Pro-Tip for Unique Reports:** Use the 'Observation Formula' for better AI results:
**[Action]** + **[Specific Tool/Component]** + **[Specific Result or Quote]**
*Example: 'Struggled with the RJ45 crimping tool at first but corrected the pin alignment manually after a second attempt.'*
""")

learning_moment = st.text_area(
    "Observation Notes", 
    placeholder="What did this specific student do, say, or struggle with during this session?",
    height=150,
    help="This is the 'flavor' that makes this report different from the others. Even 1-2 sentences here will make the AI output much more realistic."
)

# --- GENERATION LOGIC ---
if st.button("Generate & Finalize Report"):
    current_time = time.time()
    time_passed = current_time - st.session_state.last_request_time
    
    if time_passed < 10 and not dev_mode:
        st.warning(f"🕒 Rate Limit Protection: Wait {int(10 - time_passed)}s.")
    elif not selected_pcs:
        st.warning("Select at least one PC.")
    elif not target_key and not dev_mode:
        st.warning(f"Please enter the {ai_provider} API key in the sidebar.")
    else:
        st.session_state.last_request_time = current_time
        with st.spinner(f"Using {ai_provider} ({target_model}) to synthesize..."):
            
            unique_units = list(set([pc.split(' - ')[0] for pc in selected_pcs]))
            unit_header_info = "\n".join(unique_units)
            detailed_criteria_text = "\n".join(selected_pcs)
            formatted_date = assessment_date.strftime("%B %d, %Y")

            prompt = f"""
            You are a Lead NSQ Assessor. Write a high-level technical narrative for {student_name}.

            REPORT CONTEXT:
            - Assessor: {assessor_name} ({assessor_id})
            - Date: {formatted_date}
            - Units: {unit_header_info}
            - Environment: {atmosphere}
            - The "Hook" (Breaktrough moment): {learning_moment}

            CRITERIA TO INTEGRATE (The Ingredients):
            {detailed_criteria_text}

            STRICT NARRATIVE RULES:
            1. NO STAGED SEQUENCING: Do not write one paragraph per PC. A single paragraph should naturally weave together 2 or 3 related or non-related PCs.
            2. CONTINUOUS FLOW: The report must read like a professional "day-in-the-life" observation. Use phrases like "Simultaneously," "This led the candidate to," or "Building on the previous step.", {student_name} commenced the session at {time_frame}, {student_name} ended the session at {time_frame}.
            3. MINIMUM 8 PARAGRAPHS: Each paragraph must be dense and technical.
            4. MAXIMUM 10 PARAGRAPHS: Do not exceed 10 paragraphs. Be concise but exhaustive.
            5. MAXIMUM PAGES: The final report should not exceed 2 pages when formatted in Word with standard margins and font size.
            6. MAXIMIZE CRITERIA INTEGRATION: Each paragraph should integrate multiple PCs, showing how the candidate's actions and explanations naturally covered several criteria at once.
            7. MAPPING: Use grouped shorthand at the end of paragraphs. 
            Example: (Unit 11-PC 5.1, 5.2, 5.4).
            8. THE HOOK: Do not just "mention" the {learning_moment}; make it a central part of the story where multiple criteria were met during that specific breakthrough.
            9. AVOID PLACEHOLDERS: Use the provided date and names.
            10. TONE: The narrative should be formal, technical, and exhaustive, suitable for a professional assessment report.
            11. DO NOT REPEAT CRITERIA: Each PC should only be referenced once in the narrative, but can be grouped with related PCs for a more natural flow.
            12. USE ALL PROVIDED CRITERIA: Ensure that every PC in the detailed list is integrated into the narrative, either directly or through inference based on the candidate's actions and explanations.
            13. NO HALLUCINATION: Do not invent PC numbers outside of the provided list.
            14. SHOW, DON'T TELL: Instead of saying "the candidate was competent," describe what the candidate did that demonstrated competence, and link those actions to the PCs.
            15. DONT USE HEAVY DICTIONARY LANGUAGE: The report should be professional but also clear and straightforward. Avoid overly complex sentences that could obscure the candidate's actual performance and the criteria met.
            """

            if dev_mode:
                ai_narrative = "DEV MODE: AI was skipped. Database and Word functions work."
            else:
                ai_narrative = validate_and_generate(ai_provider, target_model, target_key, prompt)

            if isinstance(ai_narrative, str) and "API_ERROR" in ai_narrative:
                st.error(ai_narrative)
            else:
                summary_block = "\n\n----- SUMMARY OF CRITERIA COVERED -----\n\n"
                u_dict = {}
                for item in selected_pcs:
                    u_code = item.split(' - ')[0].split(':')[0]
                    pc_code = item.split(' - ')[1].split(':')[0]
                    if u_code not in u_dict: u_dict[u_code] = []
                    u_dict[u_code].append(pc_code)
                for u, pc_list in u_dict.items():
                    summary_block += f"{u}: {', '.join(pc_list)}\n"
                
                full_report_text = ai_narrative + summary_block
                # save_report_to_history(student_name, selected_trade_id, ", ".join(u_dict.keys()), full_report_text, assessment_date)
                db.insert_report(student_name, selected_trade_id, ", ".join(u_dict.keys()), full_report_text, assessment_date)
                st.markdown("### Preview")
                st.info(full_report_text)
                doc_bytes = export_to_word(student_name, assessment_date, time_frame, atmosphere, full_report_text, selected_pcs, assessor_name, assessor_id)
                
                c1, c2 = st.columns(2)
                with c1: st.download_button("📥 Word (.docx)", data=doc_bytes, file_name=f"NSQ_{student_name}.docx")
                with c2: st.download_button("Download Text (.txt)", full_report_text, file_name=f"{student_name}.txt")